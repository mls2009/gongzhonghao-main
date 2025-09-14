from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session
from models.database import get_db, Material, Settings
from pydantic import BaseModel
from typing import Optional
from datetime import datetime
import os
import glob
import docx
import subprocess
import platform
import traceback
import logging

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/settings", tags=["settings"])

class PathUpdate(BaseModel):
    materials_path: str

@router.post("/select-folder")
async def select_folder():
    try:
        if platform.system() == "Darwin":  # macOS
            # 使用AppleScript的文件夹选择对话框
            applescript = '''
            tell application "System Events"
                activate
                set folderPath to choose folder with prompt "选择素材库文件夹"
                return POSIX path of folderPath
            end tell
            '''
            
            result = subprocess.run(
                ["osascript", "-e", applescript],
                capture_output=True,
                text=True
            )
            
            if result.returncode == 0 and result.stdout.strip():
                folder_path = result.stdout.strip().rstrip('/')
                return {"success": True, "path": folder_path}
        
        elif platform.system() == "Windows":
            # 使用PowerShell的文件夹选择对话框
            ps_script = '''
            Add-Type -AssemblyName System.Windows.Forms
            $folderBrowser = New-Object System.Windows.Forms.FolderBrowserDialog
            $folderBrowser.Description = "选择素材库文件夹"
            $folderBrowser.RootFolder = [System.Environment+SpecialFolder]::Desktop
            $folderBrowser.ShowNewFolderButton = $true
            if ($folderBrowser.ShowDialog() -eq [System.Windows.Forms.DialogResult]::OK) {
                $folderBrowser.SelectedPath
            }
            '''
            
            result = subprocess.run(
                ["powershell", "-Command", ps_script],
                capture_output=True,
                text=True
            )
            
            if result.returncode == 0 and result.stdout.strip():
                folder_path = result.stdout.strip()
                return {"success": True, "path": folder_path}
        
        else:
            # Linux或其他系统，使用tkinter作为后备方案
            try:
                import tkinter as tk
                from tkinter import filedialog
                
                root = tk.Tk()
                root.withdraw()  # 隐藏主窗口
                folder_path = filedialog.askdirectory(title="选择素材库文件夹")
                root.destroy()
                
                if folder_path:
                    return {"success": True, "path": folder_path}
            except ImportError:
                return {"success": False, "message": "当前系统不支持图形化文件夹选择"}
        
        return {"success": False, "message": "未选择文件夹"}
            
    except Exception as e:
        print(f"Error in select_folder: {str(e)}")  # 添加日志输出
        return {"success": False, "message": f"选择文件夹时出错: {str(e)}"}

@router.post("/save-path")
async def save_path(path_update: PathUpdate, db: Session = Depends(get_db)):
    try:
        # 检查路径是否存在
        if not os.path.exists(path_update.materials_path):
            raise HTTPException(status_code=400, detail="指定的路径不存在")
        
        # 获取或创建设置记录
        settings = db.query(Settings).first()
        if not settings:
            settings = Settings(materials_path=path_update.materials_path)
            db.add(settings)
        else:
            settings.materials_path = path_update.materials_path
        
        db.commit()
        return {"success": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/scan-materials")
async def scan_materials(db: Session = Depends(get_db)):
    try:
        # 获取设置
        settings = db.query(Settings).first()
        if not settings or not settings.materials_path:
            logger.error("未设置素材库路径")
            return {"success": False, "message": "未设置素材库路径"}
            
        # 获取所有 .docx 文件
        docx_files = glob.glob(os.path.join(settings.materials_path, "**/*.docx"), recursive=True)
        logger.info(f"找到 {len(docx_files)} 个Word文档")
        
        # 获取所有在已发布列表中的素材（不管发布成功还是失败，包括隐藏的）
        published_materials = {
            material.original_title: material 
            for material in db.query(Material).filter(
                Material.status.in_(["published", "scheduled", "hidden"])
            ).all()
        }
        
        # 单独获取隐藏状态的素材，用于恢复
        hidden_materials = {
            material.original_title: material 
            for material in db.query(Material).filter(
                Material.status == "hidden"
            ).all()
        }
        
        # 清除现有的未发布素材
        db.query(Material).filter(Material.status == "unpublished").delete()
        
        # 处理每个文档
        success_count = 0
        skipped_count = 0
        restored_count = 0
        for file_path in docx_files:
            try:
                # 获取文件名（不包含扩展名）
                file_name = os.path.splitext(os.path.basename(file_path))[0]
                
                # 检查是否是隐藏状态的素材，如果是则恢复
                if file_name in hidden_materials:
                    hidden_material = hidden_materials[file_name]
                    hidden_material.status = "published"  # 恢复为已发布状态
                    logger.info(f"恢复隐藏素材: {file_name}")
                    restored_count += 1
                    continue
                
                # 如果文件已经在已发布列表中（非隐藏），跳过
                if file_name in published_materials and file_name not in hidden_materials:
                    logger.info(f"跳过已发布文件: {file_name}")
                    skipped_count += 1
                    continue
                    
                doc = docx.Document(file_path)
                
                # 计算字数
                word_count = sum(len(paragraph.text) for paragraph in doc.paragraphs)
                
                # 计算图片数量
                image_count = 0
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        image_count += 1
                
                # 读取文档内容
                content = "\n".join(paragraph.text for paragraph in doc.paragraphs)
                
                # 创建素材记录
                material = Material(
                    title=file_name,
                    original_title=file_name,
                    content=content,
                    word_count=word_count,
                    image_count=image_count,
                    status="unpublished",
                    publish_status=None,
                    publish_time=None,
                    account_id=None,
                    error_message=None,
                    schedule_time=None,
                    schedule_status=None
                )
                db.add(material)
                success_count += 1
                logger.info(f"成功处理文件: {file_path}")
                
            except Exception as doc_error:
                logger.error(f"处理文件 {file_path} 时出错: {str(doc_error)}")
                continue
        
        # 更新设置
        settings.last_scan = datetime.now()
        settings.total_files = len(docx_files)
        
        # 提交更改
        db.commit()
        
        message_parts = [f"扫描完成，共找到 {len(docx_files)} 个文档"]
        if success_count > 0:
            message_parts.append(f"处理 {success_count} 个新文档")
        if skipped_count > 0:
            message_parts.append(f"跳过 {skipped_count} 个已发布文档")
        if restored_count > 0:
            message_parts.append(f"恢复 {restored_count} 个之前清空的文档")
        
        return {
            "success": True,
            "total_files": len(docx_files),
            "processed_files": success_count,
            "skipped_files": skipped_count,
            "restored_files": restored_count,
            "message": "，".join(message_parts)
        }
        
    except Exception as e:
        logger.error(f"扫描素材库时出错: {str(e)}")
        logger.error(traceback.format_exc())
        db.rollback()
        return {"success": False, "message": f"扫描素材库时出错: {str(e)}"}

@router.get("/stats")
async def get_stats(db: Session = Depends(get_db)):
    try:
        settings = db.query(Settings).first()
        if not settings:
            return {
                "total_files": 0,
                "last_update": None,
                "materials_path": None
            }
        
        return {
            "total_files": settings.total_files,
            "last_update": settings.last_scan,
            "materials_path": settings.materials_path
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e)) 