from fastapi import FastAPI, Request, Depends, HTTPException
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from sqlalchemy.orm import Session
from models.database import get_db, Account, Browser, Material, Settings, SessionLocal, engine, Base
from routers import (
    materials_router,
    browsers_router,
    accounts_router,
    settings_router
)
from routers.xiaohongshu_materials import router as xiaohongshu_materials_router
from routers.xiaohongshu_settings import router as xiaohongshu_settings_router
from routers.template_materials import router as template_materials_router
from typing import Dict, Optional, List
from pydantic import BaseModel
import requests
import json
import logging
import traceback
import os
import glob
import docx
from datetime import datetime
import asyncio
from fastapi.background import BackgroundTasks
from scheduler.publish_scheduler import init_scheduler

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI()

# 初始化数据库表
Base.metadata.create_all(bind=engine)

# 初始化定时任务
init_scheduler()

# 注册路由
app.include_router(materials_router)
app.include_router(browsers_router)
app.include_router(accounts_router)
app.include_router(settings_router)
app.include_router(xiaohongshu_materials_router)
app.include_router(xiaohongshu_settings_router)
app.include_router(template_materials_router)

async def auto_scan_materials(db: Session):
    """应用启动时自动扫描素材库"""
    try:
        # 获取设置
        settings = db.query(Settings).first()
        if not settings or not settings.materials_path:
            logger.info("未设置素材库路径，跳过自动扫描")
            return
            
        # 获取所有 .docx 文件
        docx_files = glob.glob(os.path.join(settings.materials_path, "**/*.docx"), recursive=True)
        logger.info(f"找到 {len(docx_files)} 个Word文档")
        
        # 获取所有在已发布列表中的素材（不管发布成功还是失败）
        published_materials = {
            material.original_title: material.status 
            for material in db.query(Material).filter(
                Material.status == "published"
            ).all()
        }
        
        # 清除现有的未发布素材
        db.query(Material).filter(Material.status == "unpublished").delete()
        
        # 处理每个文档
        success_count = 0
        skipped_count = 0
        for file_path in docx_files:
            try:
                # 获取文件名（不包含扩展名）
                file_name = os.path.splitext(os.path.basename(file_path))[0]
                
                # 如果文件已经在已发布列表中，跳过
                if file_name in published_materials:
                    logger.info(f"跳过已发布文件: {file_name}")
                    skipped_count += 1
                    continue
                    
                doc = docx.Document(file_path)
                
                # 计算字数
                word_count = sum(len(paragraph.text) for paragraph in doc.paragraphs)
                
                # 计算图片数量
                image_count = 0
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        image_count += 1
                
                # 读取文档内容
                content = "\n".join(paragraph.text for paragraph in doc.paragraphs)
                
                # 创建素材记录
                material = Material(
                    title=file_name,
                    original_title=file_name,
                    content=content,
                    word_count=word_count,
                    image_count=image_count,
                    status="unpublished"
                )
                db.add(material)
                success_count += 1
                logger.info(f"成功处理文件: {file_path}")
                
            except Exception as doc_error:
                logger.error(f"处理文件 {file_path} 时出错: {str(doc_error)}")
                continue
        
        # 更新设置
        settings.last_scan = datetime.now()
        settings.total_files = len(docx_files)
        
        db.commit()
        logger.info(f"自动扫描完成，共找到 {len(docx_files)} 个文档，"
                   f"处理 {success_count} 个，"
                   f"跳过 {skipped_count} 个已发布文档")
        
    except Exception as e:
        logger.error(f"自动扫描素材库时出错: {str(e)}")
        db.rollback()

@app.on_event("startup")
async def startup_event():
    """应用启动时执行的操作"""
    db = None
    try:
        # 创建数据库会话
        db = SessionLocal()
        
        # 执行自动扫描
        await auto_scan_materials(db)
        
    except Exception as e:
        logger.error(f"启动事件出错: {str(e)}")
        logger.error(traceback.format_exc())
    finally:
        if db:
            db.close()

# BitBrowser API配置
BITBROWSER_URL = "http://127.0.0.1:54345"
BITBROWSER_HEADERS = {'Content-Type': 'application/json'}

# 设置静态文件目录
app.mount("/static", StaticFiles(directory="static"), name="static")

# 设置模板目录
templates = Jinja2Templates(directory="templates")

# Pydantic models
class AccountCreate(BaseModel):
    username: str
    password: str
    author_name: str
    account_type: str
    browser_id: str
    browser_name: str
    status: str = "active"

class AccountResponse(BaseModel):
    id: int
    author_name: str
    account_type: str

    class Config:
        orm_mode = True

class HomepageUpdate(BaseModel):
    homepage: str

@app.get("/")
async def index(request: Request):
    return templates.TemplateResponse("index.html", {
        "request": request,
        "title": "首页"
    })

@app.get("/accounts")
async def accounts(request: Request, db: Session = Depends(get_db)):
    accounts = db.query(Account).all()
    return templates.TemplateResponse("accounts.html", {
        "request": request,
        "accounts": accounts
    })

@app.post("/api/accounts")
async def create_account(account: AccountCreate, db: Session = Depends(get_db)):
    try:
        db_account = Account(**account.dict())
        db.add(db_account)
        db.commit()
        db.refresh(db_account)
        return {"success": True, "data": db_account}
    except Exception as e:
        db.rollback()
        logger.error(f"Error creating account: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

class AccountUpdate(AccountCreate):
    pass

@app.put("/api/accounts/{account_id}")
async def update_account(account_id: int, account: AccountUpdate, db: Session = Depends(get_db)):
    try:
        db_account = db.query(Account).filter(Account.id == account_id).first()
        if not db_account:
            raise HTTPException(status_code=404, detail="账号不存在")

        data = account.dict()
        for k, v in data.items():
            setattr(db_account, k, v)

        db.commit()
        db.refresh(db_account)
        return {"success": True, "data": db_account}
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"Error updating account: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/accounts/{account_id}")
async def get_account(account_id: int, db: Session = Depends(get_db)):
    account = db.query(Account).filter(Account.id == account_id).first()
    if not account:
        raise HTTPException(status_code=404, detail="账号不存在")
    return {"success": True, "data": account}

@app.delete("/api/accounts/{account_id}")
async def delete_account(account_id: int, db: Session = Depends(get_db)):
    account = db.query(Account).filter(Account.id == account_id).first()
    if not account:
        raise HTTPException(status_code=404, detail="账号不存在")
    try:
        db.delete(account)
        db.commit()
        return {"success": True, "message": "账号已删除"}
    except Exception as e:
        db.rollback()
        logger.error(f"Error deleting account: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/accounts/{account_id}/homepage")
async def update_account_homepage(
    account_id: int,
    homepage_data: HomepageUpdate,
    db: Session = Depends(get_db)
):
    try:
        account = db.query(Account).filter(Account.id == account_id).first()
        if not account:
            raise HTTPException(status_code=404, detail="账号不存在")
        
        account.homepage = homepage_data.homepage
        db.commit()
        
        return {"success": True, "message": "账号主页已更新"}
    except Exception as e:
        db.rollback()
        logger.error(f"Error updating account homepage: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/browsers/list")
async def list_browsers():
    try:
        json_data = {
            "page": 0,
            "pageSize": 100,
            "platform": "",
            "groupId": "",
            "name": "",
            "status": ""
        }
        logger.info(f"Sending request to BitBrowser API: {BITBROWSER_URL}/browser/list")
        logger.info(f"Request data: {json_data}")
        
        try:
            response = requests.post(
                f"{BITBROWSER_URL}/browser/list",
                data=json.dumps(json_data),
                headers=BITBROWSER_HEADERS,
                timeout=10
            )
            response.raise_for_status()
        except requests.exceptions.ConnectionError:
            logger.error("无法连接到BitBrowser API，请确保BitBrowser正在运行")
            raise HTTPException(status_code=503, detail="无法连接到BitBrowser，请确保BitBrowser正在运行")
        except requests.exceptions.Timeout:
            logger.error("请求BitBrowser API超时")
            raise HTTPException(status_code=504, detail="请求BitBrowser超时")
        except requests.exceptions.RequestException as e:
            logger.error(f"请求BitBrowser API时发生错误: {str(e)}")
            raise HTTPException(status_code=500, detail=f"请求BitBrowser时发生错误: {str(e)}")
        
        try:
            response_data = response.json()
        except json.JSONDecodeError:
            logger.error(f"无法解析BitBrowser API的响应: {response.text}")
            raise HTTPException(status_code=500, detail="无法解析BitBrowser的响应")
        
        logger.info(f"Response from BitBrowser API: {response_data}")
        
        if not isinstance(response_data, dict):
            raise HTTPException(status_code=500, detail="BitBrowser返回了无效的数据格式")
        
        if 'data' not in response_data or 'list' not in response_data['data']:
            raise HTTPException(status_code=500, detail="BitBrowser返回的数据缺少必要的字段")
        
        if not response_data['data']['list']:
            logger.warning("BitBrowser返回了空的浏览器列表")
        
        return response_data
    except Exception as e:
        logger.error(f"Error in list_browsers: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/browsers/{browser_id}/open")
async def open_browser(browser_id: str):
    try:
        json_data = {"id": browser_id}
        logger.info(f"Opening browser: {browser_id}")
        
        response = requests.post(
            f"{BITBROWSER_URL}/browser/open",
            data=json.dumps(json_data),
            headers=BITBROWSER_HEADERS
        )
        return response.json()
    except Exception as e:
        logger.error(f"Error in open_browser: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/browsers/create")
async def create_browser(browser_data: Dict):
    try:
        logger.info(f"Creating browser with data: {browser_data}")
        
        response = requests.post(
            f"{BITBROWSER_URL}/browser/update",
            data=json.dumps(browser_data),
            headers=BITBROWSER_HEADERS
        )
        return response.json()
    except Exception as e:
        logger.error(f"Error in create_browser: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/browsers/{browser_id}/delete")
async def delete_browser(browser_id: str):
    try:
        json_data = {"id": browser_id}
        logger.info(f"Deleting browser: {browser_id}")
        
        response = requests.post(
            f"{BITBROWSER_URL}/browser/delete",
            data=json.dumps(json_data),
            headers=BITBROWSER_HEADERS
        )
        return response.json()
    except Exception as e:
        logger.error(f"Error in delete_browser: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/settings")
async def settings_page(request: Request, db: Session = Depends(get_db)):
    try:
        # 获取设置信息
        settings = db.query(Settings).first()
        
        # 获取统计信息
        stats = {
            "total_files": settings.total_files if settings else 0,
            "last_update": settings.last_scan if settings else None
        }
        
        return templates.TemplateResponse(
            "settings.html",
            {
                "request": request,
                "title": "系统设置",
                "settings": settings,
                "stats": stats
            }
        )
    except Exception as e:
        logger.error(f"Error in settings_page: {str(e)}")
        logger.error(traceback.format_exc())
        return templates.TemplateResponse(
            "settings.html",
            {
                "request": request,
                "title": "系统设置",
                "error": str(e),
                "settings": None,
                "stats": None
            }
        )

@app.get("/materials")
async def materials_page(request: Request, db: Session = Depends(get_db)):
    try:
        logger.info("Fetching materials data...")
        
        # 获取未发布的素材
        unpublished_materials = db.query(Material).filter(Material.status == "unpublished").all()
        logger.info(f"Found {len(unpublished_materials)} unpublished materials")
        
        # 获取已发布的素材
        published_materials = db.query(Material).filter(Material.status == "published").all()
        logger.info(f"Found {len(published_materials)} published materials")
        
        # 获取所有活跃账号并序列化
        accounts = db.query(Account).filter(Account.status == "active").all()
        logger.info(f"Found {len(accounts)} active accounts")
        
        # 序列化账号数据
        serialized_accounts = []
        for account in accounts:
            serialized_accounts.append({
                "id": account.id,
                "username": account.username,
                "author_name": account.author_name,
                "account_type": account.account_type
            })
        
        # 填充账号信息
        for material in unpublished_materials + published_materials:
            if material.account_id:
                account = db.query(Account).filter(Account.id == material.account_id).first()
                if account:
                    material.author_name = account.author_name
                    material.account_type = account.account_type
        
        return templates.TemplateResponse(
            "materials.html",
            {
                "request": request,
                "title": "素材库",
                "unpublished_materials": unpublished_materials,
                "published_materials": published_materials,
                "accounts": serialized_accounts  # 使用序列化后的账号数据
            }
        )
    except Exception as e:
        logger.error(f"Error in materials_page: {str(e)}")
        logger.error(traceback.format_exc())
        return templates.TemplateResponse(
            "materials.html",
            {
                "request": request,
                "title": "素材库",
                "error": str(e),
                "unpublished_materials": [],
                "published_materials": [],
                "accounts": []
            }
        ) 

@app.get("/xiaohongshu-materials")
async def xiaohongshu_materials(request: Request, db: Session = Depends(get_db)):
    try:
        logger.info("Loading xiaohongshu materials page...")
        
        # 获取小红书账号列表
        accounts = db.query(Account).filter(Account.account_type == "xiaohongshu", Account.status == "active").all()
        logger.info(f"Found {len(accounts)} active xiaohongshu accounts")
        
        # 序列化账号数据
        serialized_accounts = []
        for account in accounts:
            serialized_accounts.append({
                "id": account.id,
                "username": account.username,
                "author_name": account.author_name,
                "account_type": account.account_type,
                "status": account.status
            })
        
        return templates.TemplateResponse("xiaohongshu_materials.html", {
            "request": request,
            "title": "小红书素材",
            "accounts": serialized_accounts,
            "error": None
        })
    except Exception as e:
        logger.error(f"Error loading xiaohongshu materials page: {str(e)}")
        return templates.TemplateResponse("xiaohongshu_materials.html", {
            "request": request,
            "title": "小红书素材",
            "accounts": [],
            "error": str(e)
        })

@app.get("/template-materials")
async def template_materials(request: Request):
    return templates.TemplateResponse("template_materials.html", {
        "request": request,
        "title": "素材模板"
    }) 
